"""
RAG Ingestion — parses PDF, DOCX, TXT files into LangChain Document chunks.
Attaches source metadata for citation tracking.
"""
from __future__ import annotations

from pathlib import Path
from typing import List

try:
    from langchain_core.documents import Document
except ImportError:
    from langchain.schema import Document
try:
    from langchain_text_splitters.character import RecursiveCharacterTextSplitter
except ImportError:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
from loguru import logger

from config.settings import CHUNK_SIZE, CHUNK_OVERLAP


def _read_pdf(file_path: str) -> List[dict]:
    """Extract text from PDF page by page using PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        pages = []
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                pages.append({"text": text, "page": page_num})
        doc.close()
        return pages
    except Exception as e:
        logger.error(f"[Ingestion] PDF read error for '{file_path}': {e}")
        return []


def _read_docx(file_path: str) -> List[dict]:
    """Extract text from DOCX files using python-docx."""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [{"text": full_text, "page": 1}]
    except Exception as e:
        logger.error(f"[Ingestion] DOCX read error for '{file_path}': {e}")
        return []


def _read_txt(file_path: str) -> List[dict]:
    """Read plain text files with encoding detection."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
            if text:
                return [{"text": text, "page": 1}]
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"[Ingestion] TXT read error for '{file_path}': {e}")
            return []
    logger.warning(f"[Ingestion] Could not decode '{file_path}' with any encoding.")
    return []


def ingest_file(file_path: str) -> List[Document]:
    """
    Parse a single file and return a list of LangChain Document chunks.
    Supports PDF, DOCX, TXT.
    """
    path = Path(file_path)
    if not path.exists():
        logger.error(f"[Ingestion] File not found: {file_path}")
        return []

    ext = path.suffix.lower()
    filename = path.name

    if ext == ".pdf":
        raw_pages = _read_pdf(file_path)
        file_type = "pdf"
    elif ext in (".docx", ".doc"):
        raw_pages = _read_docx(file_path)
        file_type = "docx"
    elif ext == ".txt":
        raw_pages = _read_txt(file_path)
        file_type = "txt"
    else:
        logger.warning(f"[Ingestion] Unsupported file type: {ext}")
        return []

    if not raw_pages:
        logger.warning(f"[Ingestion] No text extracted from '{filename}'")
        return []

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    documents = []
    for page_data in raw_pages:
        chunks = splitter.split_text(page_data["text"])
        for chunk_idx, chunk_text in enumerate(chunks):
            if not chunk_text.strip():
                continue
            doc = Document(
                page_content=chunk_text,
                metadata={
                    "source": filename,
                    "file_path": str(path.resolve()),
                    "page": page_data["page"],
                    "chunk_index": chunk_idx,
                    "file_type": file_type,
                },
            )
            documents.append(doc)

    logger.info(
        f"[Ingestion] '{filename}' → {len(raw_pages)} page(s), {len(documents)} chunks"
    )
    return documents


def ingest_directory(directory: str) -> List[Document]:
    """Ingest all supported files from a directory."""
    dir_path = Path(directory)
    if not dir_path.is_dir():
        logger.error(f"[Ingestion] Directory not found: {directory}")
        return []

    supported = {".pdf", ".docx", ".doc", ".txt"}
    all_docs = []
    for file in sorted(dir_path.iterdir()):
        if file.suffix.lower() in supported:
            docs = ingest_file(str(file))
            all_docs.extend(docs)

    logger.info(f"[Ingestion] Directory '{directory}' → {len(all_docs)} total chunks")
    return all_docs
