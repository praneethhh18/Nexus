"""
Document generation module — render Word (.docx) and PDF business documents
from templates (Proposal, Statement of Work, Contract, Offer Letter) with
variable substitution.

All generated documents are stored under outputs/documents/<business_id>/
and tracked in the nexus_documents table so the UI can list them.
"""
from __future__ import annotations

import json
import re
import sqlite3  # sqlite3.Row sentinel — works on Postgres via config.db
import uuid
from datetime import datetime, date
from pathlib import Path
from typing import List, Dict, Any

from fastapi import HTTPException
from loguru import logger

from config.db import get_conn, list_columns
from config.settings import OUTPUTS_DIR

DOCS_TABLE = "nexus_documents"
DOCS_DIR = Path(OUTPUTS_DIR) / "documents"


# ── Templates ────────────────────────────────────────────────────────────────
# A template is a plain dict with variable placeholders like {{customer_name}}.
# The UI shows the list via list_templates() and the variables it expects.
TEMPLATES: Dict[str, Dict[str, Any]] = {
    "proposal": {
        "name": "Business Proposal",
        "description": "Cover letter + scope + pricing for a new prospect",
        "variables": [
            "customer_name", "customer_company", "project_name",
            "project_summary", "scope_bullets", "timeline",
            "total_amount", "currency", "valid_until",
            "sender_name", "sender_title",
        ],
        "sections": [
            ("heading", "Business Proposal: {{project_name}}"),
            ("paragraph", "Prepared for: {{customer_name}}{{#customer_company}}, {{customer_company}}{{/customer_company}}"),
            ("paragraph", "Date: {{today}}"),
            ("paragraph", "Valid until: {{valid_until}}"),
            ("heading2", "Executive Summary"),
            ("paragraph", "{{project_summary}}"),
            ("heading2", "Scope of Work"),
            ("bullets", "{{scope_bullets}}"),
            ("heading2", "Timeline"),
            ("paragraph", "{{timeline}}"),
            ("heading2", "Investment"),
            ("paragraph", "Total: {{total_amount}} {{currency}}"),
            ("paragraph", "Please confirm acceptance by signing below or replying to this document."),
            ("signoff", "{{sender_name}}\n{{sender_title}}"),
        ],
    },
    "sow": {
        "name": "Statement of Work",
        "description": "Detailed scope, deliverables, and acceptance criteria",
        "variables": [
            "customer_name", "customer_company", "project_name",
            "objectives", "deliverables", "milestones",
            "acceptance_criteria", "total_amount", "currency",
            "start_date", "end_date", "sender_name", "sender_title",
        ],
        "sections": [
            ("heading", "Statement of Work: {{project_name}}"),
            ("paragraph", "Between: {{sender_name}} and {{customer_name}}{{#customer_company}}, {{customer_company}}{{/customer_company}}"),
            ("paragraph", "Effective: {{start_date}} — {{end_date}}"),
            ("heading2", "Objectives"),
            ("paragraph", "{{objectives}}"),
            ("heading2", "Deliverables"),
            ("bullets", "{{deliverables}}"),
            ("heading2", "Milestones"),
            ("bullets", "{{milestones}}"),
            ("heading2", "Acceptance Criteria"),
            ("paragraph", "{{acceptance_criteria}}"),
            ("heading2", "Fees"),
            ("paragraph", "Total: {{total_amount}} {{currency}}"),
            ("signoff", "Authorized by {{sender_name}}, {{sender_title}}"),
        ],
    },
    "contract": {
        "name": "Service Agreement",
        "description": "Short-form service contract with terms and signatures",
        "variables": [
            "customer_name", "customer_company", "sender_company",
            "service_description", "fee_amount", "currency",
            "payment_terms", "term_length",
            "sender_name", "sender_title",
        ],
        "sections": [
            ("heading", "Service Agreement"),
            ("paragraph", "This agreement is made on {{today}} between {{sender_company}} ('Provider') and {{customer_company}} ('Client'), represented by {{customer_name}}."),
            ("heading2", "1. Services"),
            ("paragraph", "Provider will deliver: {{service_description}}"),
            ("heading2", "2. Fees and Payment"),
            ("paragraph", "Client agrees to pay {{fee_amount}} {{currency}}. Payment terms: {{payment_terms}}."),
            ("heading2", "3. Term"),
            ("paragraph", "This agreement is valid for {{term_length}} from the effective date above."),
            ("heading2", "4. Confidentiality"),
            ("paragraph", "Both parties agree to keep each other's confidential information private during and after the term of this agreement."),
            ("heading2", "5. Termination"),
            ("paragraph", "Either party may terminate this agreement with 30 days written notice. Outstanding fees remain payable."),
            ("signoff", "Provider: {{sender_name}}, {{sender_title}}\n\nClient: {{customer_name}}"),
        ],
    },
    "offer_letter": {
        "name": "Offer Letter",
        "description": "Employment / contractor offer letter",
        "variables": [
            "candidate_name", "role_title", "start_date",
            "compensation", "currency", "reporting_to",
            "work_location", "sender_name", "sender_company",
        ],
        "sections": [
            ("heading", "Offer of Engagement"),
            ("paragraph", "Dear {{candidate_name}},"),
            ("paragraph", "We are pleased to offer you the position of {{role_title}} at {{sender_company}}, with a proposed start date of {{start_date}}."),
            ("heading2", "Compensation"),
            ("paragraph", "{{compensation}} {{currency}}"),
            ("heading2", "Reporting and Location"),
            ("paragraph", "You will report to {{reporting_to}}. Primary work location: {{work_location}}."),
            ("paragraph", "To accept this offer, please sign and return this letter. We look forward to working with you."),
            ("signoff", "{{sender_name}}\n{{sender_company}}"),
        ],
    },
}


# ── Template substitution ────────────────────────────────────────────────────
_SIMPLE_VAR = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")
_SECTION_VAR = re.compile(
    r"\{\{#\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}(.+?)\{\{/\s*\1\s*\}\}",
    re.DOTALL,
)


def _resolve(text: str, vars_map: Dict[str, Any]) -> str:
    """Expand {{var}} and conditional {{#var}}…{{/var}} placeholders."""
    # Conditionals: keep block only if the var is truthy
    def _cond(m: re.Match) -> str:
        key, inner = m.group(1), m.group(2)
        return inner if vars_map.get(key) else ""
    text = _SECTION_VAR.sub(_cond, text)

    # Simple variables
    def _var(m: re.Match) -> str:
        key = m.group(1)
        v = vars_map.get(key, "")
        return str(v) if v is not None else ""
    return _SIMPLE_VAR.sub(_var, text)


def _lines_to_bullets(value: Any) -> List[str]:
    """Normalise a scope/deliverable/milestone field into bullet list."""
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]
    if not value:
        return []
    s = str(value)
    # Split on newlines or semicolons, keep non-empty
    parts = re.split(r"[\n;]+", s)
    return [p.strip(" -•*") for p in parts if p.strip()]


# ── Database ─────────────────────────────────────────────────────────────────
def _get_conn():
    conn = get_conn()
    conn.execute(f"""
    CREATE TABLE IF NOT EXISTS {DOCS_TABLE} (
        id TEXT PRIMARY KEY,
        business_id TEXT NOT NULL,
        template_key TEXT NOT NULL,
        title TEXT NOT NULL,
        format TEXT NOT NULL,
        file_path TEXT NOT NULL,
        variables TEXT DEFAULT '{{}}',
        created_at TEXT,
        created_by TEXT
    )""")
    conn.execute(f"CREATE INDEX IF NOT EXISTS idx_docs_biz ON {DOCS_TABLE}(business_id, created_at)")
    # Additive migration for RAG collections + document expiry — safe to re-run.
    for col, decl in [
        ("collection_id", "TEXT"),
        ("expires_at",    "TEXT"),
    ]:
        existing = list_columns(conn, DOCS_TABLE)
        if col not in existing:
            conn.execute(f"ALTER TABLE {DOCS_TABLE} ADD COLUMN {col} {decl}")
    conn.commit()
    return conn


def _business_dir(business_id: str) -> Path:
    # business_id is validated before we get here, but sanitise anyway
    safe = re.sub(r"[^A-Za-z0-9_-]", "_", business_id)[:40]
    d = DOCS_DIR / safe
    d.mkdir(parents=True, exist_ok=True)
    return d


# ── Template listing ─────────────────────────────────────────────────────────
def list_templates() -> List[Dict[str, Any]]:
    return [
        {
            "key": key,
            "name": t["name"],
            "description": t["description"],
            "variables": t["variables"],
        }
        for key, t in TEMPLATES.items()
    ]


def get_template(template_key: str) -> Dict[str, Any]:
    if template_key not in TEMPLATES:
        raise HTTPException(404, f"Unknown template: {template_key}")
    return TEMPLATES[template_key]


# ── Generation ───────────────────────────────────────────────────────────────
def _render_docx(template_key: str, vars_map: Dict[str, Any], out_path: Path) -> None:
    """Render the template as a Word document using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tmpl = TEMPLATES[template_key]
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    accent = RGBColor(0x1E, 0x3A, 0x5F)

    for kind, raw in tmpl["sections"]:
        text = _resolve(raw, vars_map)
        if kind == "heading":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = accent
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()  # spacing
        elif kind == "heading2":
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = accent
        elif kind == "paragraph":
            doc.add_paragraph(text)
        elif kind == "bullets":
            for line in _lines_to_bullets(text):
                doc.add_paragraph(line, style="List Bullet")
        elif kind == "signoff":
            doc.add_paragraph()
            for line in text.splitlines():
                doc.add_paragraph(line)

    doc.save(str(out_path))


def _render_pdf(template_key: str, vars_map: Dict[str, Any], out_path: Path) -> None:
    """Render the template as a PDF using ReportLab."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem, HRFlowable,
    )

    ACCENT = colors.HexColor("#1e3a5f")
    tmpl = TEMPLATES[template_key]

    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Title"], textColor=ACCENT, fontSize=22, alignment=0, spaceAfter=14)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=ACCENT, fontSize=13, spaceBefore=10, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=11, textColor=colors.HexColor("#0f172a"), leading=15)

    pdf = SimpleDocTemplate(
        str(out_path), pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2 * cm, rightMargin=2 * cm,
    )
    story = []

    for kind, raw in tmpl["sections"]:
        text = _resolve(raw, vars_map).replace("\n", "<br/>")
        if kind == "heading":
            story.append(Paragraph(text, h1))
            story.append(HRFlowable(color=ACCENT, thickness=1))
            story.append(Spacer(1, 0.3 * cm))
        elif kind == "heading2":
            story.append(Paragraph(text, h2))
        elif kind == "paragraph":
            story.append(Paragraph(text, body))
            story.append(Spacer(1, 0.15 * cm))
        elif kind == "bullets":
            items = _lines_to_bullets(text)
            if items:
                story.append(ListFlowable(
                    [ListItem(Paragraph(it, body)) for it in items],
                    bulletType="bullet", start="circle",
                ))
                story.append(Spacer(1, 0.15 * cm))
        elif kind == "signoff":
            story.append(Spacer(1, 0.8 * cm))
            for line in text.splitlines():
                story.append(Paragraph(line, body))

    pdf.build(story)


def generate_document(
    business_id: str,
    user_id: str,
    template_key: str,
    title: str,
    variables: Dict[str, Any],
    fmt: str = "docx",
) -> Dict[str, Any]:
    if template_key not in TEMPLATES:
        raise HTTPException(404, f"Unknown template: {template_key}")
    if fmt not in ("docx", "pdf"):
        raise HTTPException(400, "format must be 'docx' or 'pdf'")

    title = (title or "").strip()
    if not title:
        raise HTTPException(400, "Document title is required")
    if len(title) > 200:
        raise HTTPException(400, "Title too long (max 200)")

    # Add implicit helpers
    vars_map = dict(variables or {})
    vars_map.setdefault("today", date.today().strftime("%B %d, %Y"))
    vars_map.setdefault("currency", "USD")

    safe_title = re.sub(r"[^A-Za-z0-9_-]+", "_", title)[:80] or "document"
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"{safe_title}_{stamp}.{fmt}"
    out_dir = _business_dir(business_id)
    out_path = out_dir / filename

    if fmt == "docx":
        _render_docx(template_key, vars_map, out_path)
    else:
        _render_pdf(template_key, vars_map, out_path)

    doc_id = f"doc-{uuid.uuid4().hex[:10]}"
    now = datetime.now().isoformat()
    conn = _get_conn()
    try:
        conn.execute(
            f"INSERT INTO {DOCS_TABLE} (id, business_id, template_key, title, format, "
            f"file_path, variables, created_at, created_by) VALUES (?,?,?,?,?,?,?,?,?)",
            (doc_id, business_id, template_key, title, fmt,
             str(out_path), json.dumps(vars_map, default=str), now, user_id),
        )
        conn.commit()
    finally:
        conn.close()

    logger.info(f"[Documents] Generated {template_key} '{title}' → {out_path}")
    return {
        "id": doc_id,
        "template_key": template_key,
        "title": title,
        "format": fmt,
        "file_path": str(out_path),
        "filename": filename,
        "created_at": now,
    }


def list_documents(business_id: str, limit: int = 100) -> List[Dict[str, Any]]:
    conn = _get_conn()
    conn.row_factory = sqlite3.Row
    try:
        rows = conn.execute(
            f"SELECT * FROM {DOCS_TABLE} WHERE business_id = ? ORDER BY created_at DESC LIMIT ?",
            (business_id, limit),
        ).fetchall()
    finally:
        conn.close()
    result = []
    for r in rows:
        d = dict(r)
        d["filename"] = Path(d["file_path"]).name
        d["variables"] = json.loads(d.get("variables") or "{}")
        result.append(d)
    return result


def get_document(business_id: str, document_id: str) -> Dict[str, Any]:
    conn = _get_conn()
    conn.row_factory = sqlite3.Row
    try:
        row = conn.execute(
            f"SELECT * FROM {DOCS_TABLE} WHERE id = ? AND business_id = ?",
            (document_id, business_id),
        ).fetchone()
    finally:
        conn.close()
    if not row:
        raise HTTPException(404, "Document not found")
    d = dict(row)
    d["filename"] = Path(d["file_path"]).name
    d["variables"] = json.loads(d.get("variables") or "{}")
    return d


def delete_document(business_id: str, document_id: str) -> None:
    doc = get_document(business_id, document_id)
    # Only remove files inside the managed documents dir
    try:
        p = Path(doc["file_path"])
        if p.is_file() and DOCS_DIR in p.parents:
            p.unlink(missing_ok=True)
    except Exception:
        pass
    conn = _get_conn()
    try:
        conn.execute(
            f"DELETE FROM {DOCS_TABLE} WHERE id = ? AND business_id = ?",
            (document_id, business_id),
        )
        conn.commit()
    finally:
        conn.close()
